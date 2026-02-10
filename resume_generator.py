import os
import re
from typing import Dict
from docx import Document
from docx.shared import Pt
import google.generativeai as genai
from dotenv import load_dotenv

load_dotenv()

API_KEY = os.getenv("GOOGLE_API_KEY")
if API_KEY:
    try:
        genai.configure(api_key=API_KEY)
    except Exception:
        pass

MODEL_NAME = "models/gemini-flash-latest"
model = genai.GenerativeModel(MODEL_NAME)


def _sanitize_field(value: str) -> str:
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value)).strip()


def generate_resume(data: Dict) -> str:
    """Generate resume text using Gemini. Raises RuntimeError on failure."""
    if not API_KEY:
        raise RuntimeError("Missing GOOGLE_API_KEY environment variable")

    # load prompt template if available
    prompt_template = None
    try:
        from pathlib import Path

        p = Path(__file__).parent / "prompts" / "resume_prompt.txt"
        if p.exists():
            prompt_template = p.read_text()
    except Exception:
        prompt_template = None

    name = _sanitize_field(data.get("name", ""))
    email = _sanitize_field(data.get("email", ""))
    phone = _sanitize_field(data.get("phone", ""))
    role = _sanitize_field(data.get("role", ""))
    jd = _sanitize_field(data.get("jd", ""))

    if prompt_template:
        prompt = prompt_template.format(role=role, name=name, email=email, phone=phone, content=data.get("content", ""), jd=jd)
    else:
        prompt = f"You are an ATS optimization expert. Create an ATS-friendly resume for the role of {role}. Candidate: Name: {name} Email: {email} Phone: {phone} Job Description: {jd}\n"

    try:
        response = model.generate_content(prompt)
        resume_text = response.text
        return str(resume_text)
    except Exception as exc:
        # propagate the exception as a RuntimeError so callers can decide to fallback
        raise RuntimeError(f"Failed to generate resume: {exc}") from exc


def generate_local_resume(data: Dict) -> str:
    """Local fallback resume generator that formats a plain-text resume from the provided fields.

    This is a lightweight offline fallback to use when the external model is unavailable (quota,
    network, etc.). It produces a short, ATS-friendly plain text resume using the input fields.
    """
    name = _sanitize_field(data.get("name", ""))
    email = _sanitize_field(data.get("email", ""))
    phone = _sanitize_field(data.get("phone", ""))
    role = _sanitize_field(data.get("role", ""))
    jd = _sanitize_field(data.get("jd", ""))
    content = _sanitize_field(data.get("content", ""))

    lines = []
    # Header
    if name:
        lines.append(name.upper())
    contact = []
    if email:
        contact.append(email)
    if phone:
        contact.append(phone)
    if contact:
        lines.append(' | '.join(contact))

    lines.append('')

    # Summary
    summary = f"Experienced candidate targeting {role}."
    if jd:
        # try to pull a few keywords from the JD (simple token selection)
        jd_tokens = re.findall(r"\b[a-zA-Z0-9_+#\-.]{2,}\b", jd)
        top = []
        seen = set()
        for t in jd_tokens:
            tl = t.lower()
            if tl in seen:
                continue
            if tl.isdigit():
                continue
            if len(tl) <= 2:
                continue
            seen.add(tl)
            top.append(t)
            if len(top) >= 6:
                break
        if top:
            summary += ' Key skills: ' + ', '.join(top[:6]) + '.'

    lines.append('SUMMARY')
    lines.append('- ' + summary)
    lines.append('')

    # Skills: combine obvious tokens from JD and short picks from content
    skills = []
    if jd:
        skills += top
    if content:
        content_tokens = re.findall(r"\b[A-Za-z0-9#+\-\.]{2,}\b", content)
        for t in content_tokens:
            if len(skills) >= 8:
                break
            if t not in skills:
                skills.append(t)
    skills = skills[:8]
    lines.append('SKILLS')
    if skills:
        for s in skills:
            lines.append(f"- {s}")
    else:
        lines.append("- N/A")

    lines.append('')

    # Projects / Experience: take up to 3 paragraphs from content separated by double newlines
    lines.append('PROJECTS')
    if content:
        parts = [p.strip() for p in re.split(r"\n{2,}", content) if p.strip()]
        for p in parts[:3]:
            # use first line as title if long
            sublines = p.split('\n')
            title = sublines[0][:60]
            lines.append(f"{title}")
            # one bullet summarizing
            lines.append(f"- {(' '.join(sublines[1:]) or sublines[0])[:160]}")
    else:
        lines.append('- N/A')

    lines.append('')
    lines.append('EDUCATION')
    lines.append('- N/A')

    return '\n'.join(lines)


def clean_text(text: str) -> str:
    if text is None:
        return ""
    text = text.replace("\x00", "")
    text = text.replace("\r\n", "\n")
    text = text.replace("\r", "\n")

    text = text.replace("•", "-")
    text = text.replace("–", "-")
    text = text.replace("—", "-")

    text = re.sub(r"(?m)^[\s]*[\*\+]\s+", "- ", text)

    text = re.sub(r"(?m)^\s*#+\s*", "", text)

    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"__(.*?)__", r"\1", text)
    text = re.sub(r"_(.*?)_", r"\1", text)

    text = re.sub(r"`(.*?)`", r"\1", text)

    text = re.sub(r"\[(.*?)\]\(.*?\)", r"\1", text)

    text = re.sub(r"\n{3,}", "\n\n", text)

    return text


def save_docx(resume_text: str, filename: str = "resume.docx") -> str:
    """Save plain resume text into a .docx file. Uses simple heuristics for headings and bullets."""
    doc = Document()
    resume_text = clean_text(resume_text)

    for line in resume_text.split("\n"):
        line = line.strip()
        if not line:
            continue

        if line.startswith("- ") or line.startswith("-") or line.startswith("* ") or line.startswith("+ "):
            # treat as bullet list (strip leading bullet char and whitespace)
            bullet_text = re.sub(r"^[\-\*\+]\s*", "", line)
            p = doc.add_paragraph(bullet_text, style="List Bullet")
            # leave default font size
        else:
            # detect common headings (case-insensitive)
            if re.match(r"^(summary|skills|projects|education)\b", line, re.I) or re.match(r"^[A-Z\s]{2,}$", line):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(12)
            else:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.size = Pt(11)

    doc.save(filename)
    return filename
